"""
Genera Documento_Tesis.docx desde Documento_Tesis.md usando pandoc.

Pasos:
  1. Crea 'reference_portada.docx' basado en la plantilla, añadiendo el
     estilo 'Portada-Centrado' (centrado, sin salto automático de página).
    2. Ejecuta pandoc con ese reference-doc para generar el .docx base.
    3. Postprocesa el .docx para aplicar:
         - índice automático con formato Word,
         - índice de tablas e índice de figuras,
         - saltos de página en secciones preliminares,
         - numeración de títulos de nivel 2 y 3 por capítulo.

Uso:
  python generar_tesis.py
"""

# pyright: reportPrivateUsage=false
# pylint: disable=protected-access

import math
import re
import locale
import subprocess
import sys
import tempfile
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

BASE = Path(__file__).parent
PLANTILLA = BASE / "plantilla_estilos.docx"
CARATULA_MANUAL = BASE / "caratula.docx"
CARATULA_MANUAL_ALT = BASE / "caratua.docx"
REFERENCE = BASE / "document_reference.docx"
ENTRADA = BASE / "Documento_Tesis.md"
SALIDA = BASE / "Documento_Tesis_salida.docx"
SALIDA_ALT = BASE / "Documento_Tesis_salida_nueva.docx"
CUERPO_TEMP = BASE / "_cuerpo_tesis_temp.docx"


def preparar_reference_doc():
    """Valida que exista la plantilla de estilos y la usa como reference-doc."""
    if not PLANTILLA.exists():
        raise FileNotFoundError(f"No se encontró la plantilla de estilos: {PLANTILLA.name}")
    if PLANTILLA != REFERENCE:
        Document(str(PLANTILLA)).save(str(REFERENCE))
        print(f"  Reference doc guardado: {REFERENCE.name}")
    else:
        print(f"  Se usará reference doc: {REFERENCE.name}")


def ejecutar_pandoc(origen_md, destino, reference_path):
    """Llama a pandoc para convertir el markdown al docx."""
    cmd = [
        "pandoc",
        str(origen_md),
        f"--reference-doc={reference_path}",
        f"--resource-path={BASE}",
        "--from", "markdown+fenced_divs+raw_tex+pipe_tables+table_captions",
        "--to", "docx",
        "-o", str(destino),
    ]
    print(f"\nEjecutando: {' '.join(cmd)}\n")
    resultado = subprocess.run(cmd, capture_output=True, text=False, check=False)
    preferida = locale.getpreferredencoding(False) or "utf-8"
    try:
        stdout = resultado.stdout.decode("utf-8")
    except UnicodeDecodeError:
        stdout = resultado.stdout.decode(preferida, errors="replace")
    try:
        stderr = resultado.stderr.decode("utf-8")
    except UnicodeDecodeError:
        stderr = resultado.stderr.decode(preferida, errors="replace")

    if resultado.returncode != 0:
        print("ERROR de pandoc:")
        print(stderr)
        return False
    else:
        if stderr:
            print("Advertencias pandoc:", stderr)
        if stdout:
            print(stdout)
        print(f"\nDocumento generado correctamente: {destino.name}")
        return True


def insertar_parrafo_despues(parrafo, texto=""):
    """Inserta un párrafo nuevo justo después del párrafo recibido."""
    nuevo_p = OxmlElement("w:p")
    parrafo._p.addnext(nuevo_p)  # pyright: ignore[reportPrivateUsage]
    nuevo = parrafo._parent.add_paragraph()  # pyright: ignore[reportPrivateUsage]
    nuevo._p.getparent().remove(nuevo._p)  # pyright: ignore[reportPrivateUsage]
    nuevo._p = nuevo_p  # pyright: ignore[reportPrivateUsage]
    if texto:
        nuevo.add_run(texto)
    return nuevo


# ---------------------------------------------------------------------------
# Secciones a excluir del DOCX generado.
# Cada entrada es el texto exacto del encabezado de nivel 1 (# Título).
# Se excluye ese encabezado y todo su contenido hasta el siguiente H1 (o fin).
# Modificar esta lista para incluir o excluir secciones según sea necesario.
# ---------------------------------------------------------------------------
SECCIONES_EXCLUIDAS: list[str] = [
    "Resumen",
    "Abstract",
    "Introducción",
    "Anexos",
]


def _filtrar_secciones_excluidas(texto_md: str) -> str:
    """Elimina del markdown las secciones de nivel 1 listadas en SECCIONES_EXCLUIDAS.

    Conserva el resto del contenido intacto, incluyendo saltos de página
    (\\newpage) que preceden a secciones incluidas.
    """
    if not SECCIONES_EXCLUIDAS:
        return texto_md

    lineas = texto_md.splitlines(keepends=True)
    resultado: list[str] = []
    omitiendo = False

    for linea in lineas:
        stripped = linea.strip()
        # Detectar encabezados de nivel 1
        if stripped.startswith("# ") and not stripped.startswith("## "):
            titulo = stripped[2:].strip()
            if titulo in SECCIONES_EXCLUIDAS:
                omitiendo = True
                # También eliminar el \newpage que precede a esta sección
                # si el último elemento añadido era solo ese comando
                while resultado and resultado[-1].strip() in ("\\newpage", ""):
                    resultado.pop()
                continue
            else:
                if omitiendo:
                    # Venía de una sección omitida: reinsertar salto de página
                    resultado.append("\\newpage\n\n")
                omitiendo = False

        if not omitiendo:
            resultado.append(linea)

    return "".join(resultado)


def extraer_markdown_cuerpo(origen_md):
    """Extrae el cuerpo del markdown (sin portada) y devuelve ruta temporal."""
    texto = origen_md.read_text(encoding="utf-8")
    idx = texto.find("\\newpage")
    if idx == -1:
        cuerpo = texto
    else:
        cuerpo = texto[idx + len("\\newpage"):].lstrip("\r\n")

    cuerpo = _filtrar_secciones_excluidas(cuerpo)

    with tempfile.NamedTemporaryFile(prefix="cuerpo_tesis_", suffix=".md", delete=False, encoding="utf-8", mode="w") as tmp:
        tmp.write(cuerpo)
        return Path(tmp.name)


def combinar_caratula_y_cuerpo(caratula_path, cuerpo_path, salida_path):
    """Combina la carátula manual (.docx) con el cuerpo generado por pandoc."""
    doc_caratula = Document(str(caratula_path))
    doc_cuerpo = Document(str(cuerpo_path))

    # --- Copiar estilos faltantes (ImageCaption, CaptionedFigure, etc.) ---
    caratula_styles_elem = doc_caratula.styles.element
    cuerpo_styles_elem = doc_cuerpo.styles.element
    ids_existentes = {
        s.get(qn("w:styleId"))
        for s in caratula_styles_elem.iterchildren()
        if s.tag == qn("w:style")
    }
    for style_elem in cuerpo_styles_elem.iterchildren():
        if style_elem.tag == qn("w:style"):
            sid = style_elem.get(qn("w:styleId"))
            if sid and sid not in ids_existentes:
                caratula_styles_elem.append(deepcopy(style_elem))

    # --- Copiar relaciones (imágenes, hipervínculos) del cuerpo al doc destino ---
    cuerpo_part = doc_cuerpo.part
    caratula_part = doc_caratula.part
    rid_map: dict[str, str] = {}  # rId viejo → rId nuevo
    for rel in cuerpo_part.rels.values():
        if "image" in rel.reltype:
            new_rid = caratula_part.relate_to(rel.target_part, rel.reltype)
            rid_map[rel.rId] = new_rid
        elif rel.is_external:
            new_rid = caratula_part.relate_to(
                rel.target_ref, rel.reltype, is_external=True
            )
            rid_map[rel.rId] = new_rid

    # --- Copiar cuerpo XML remapeando rIds ---
    body_caratula = doc_caratula.element.body
    body_cuerpo = doc_cuerpo.element.body

    indice_insercion = len(body_caratula)
    sect_pr_caratula = body_caratula.sectPr
    if sect_pr_caratula is not None:
        indice_insercion -= 1

    _R_EMBED = qn("r:embed")
    _R_LINK = qn("r:link")
    _R_ID = qn("r:id")

    for child in body_cuerpo.iterchildren():
        if child.tag == qn("w:sectPr"):
            continue
        copied = deepcopy(child)
        # Remapear rIds de imágenes e hipervínculos en el XML copiado
        for elem in copied.iter():
            for attr in (_R_EMBED, _R_LINK, _R_ID):
                old = elem.get(attr)
                if old and old in rid_map:
                    elem.set(attr, rid_map[old])
        body_caratula.insert(indice_insercion, copied)
        indice_insercion += 1

    doc_caratula.save(str(salida_path))


def insertar_campo_word(parrafo, instruccion, texto_placeholder="Actualice campos (F9)"):
    """Inserta un campo de Word (TOC, lista de tablas, lista de figuras)."""
    run = parrafo.add_run()
    r = run._r  # pyright: ignore[reportPrivateUsage]

    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instruccion

    fld_char_sep = OxmlElement("w:fldChar")
    fld_char_sep.set(qn("w:fldCharType"), "separate")

    text = OxmlElement("w:t")
    text.text = texto_placeholder

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    r.append(fld_char_begin)
    r.append(instr_text)
    r.append(fld_char_sep)
    r.append(text)
    r.append(fld_char_end)


def activar_actualizacion_campos(doc):
    """Hace que Word actualice TOC/campos al abrir el documento."""
    settings = doc.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def aplicar_saltos_pagina_encabezados(doc):
    """Aplica salto de página antes de los encabezados clave."""
    objetivos = {
        "Índice",
        "Índice de tablas",
        "Índice de figuras",
        "Resumen",
        "Abstract",
        "Introducción",
        "Referencias",
        "Anexos",
    }
    for p in doc.paragraphs:
        texto = p.text.strip()
        estilo = p.style.name if p.style is not None else ""

        if texto in objetivos:
            p.paragraph_format.page_break_before = True
        elif estilo == "Heading 1" and texto.startswith("Capítulo "):
            p.paragraph_format.page_break_before = True


def aplicar_numeracion_titulos(doc):
    """Numera Heading 2/3 por capítulo para aproximar el formato solicitado."""
    capitulo_actual = 0
    nivel2 = 0
    nivel3 = 0
    dentro_de_capitulo = False

    for p in doc.paragraphs:
        estilo = p.style.name if p.style is not None else ""
        texto = p.text.strip()

        if estilo == "Heading 1":
            if texto.startswith("Capítulo "):
                capitulo_actual += 1
                nivel2 = 0
                nivel3 = 0
                dentro_de_capitulo = True
                if ": " in texto:
                    p.text = texto.replace(": ", ". ", 1)
            else:
                dentro_de_capitulo = False

        elif dentro_de_capitulo and estilo == "Heading 2":
            nivel2 += 1
            nivel3 = 0
            limpio = re.sub(r"^\d+\.\d+\.?\s*(->\s*)?", "", texto)
            limpio = re.sub(r"^-\s*", "", limpio)
            p.text = f"{capitulo_actual}.{nivel2}. {limpio}"

        elif dentro_de_capitulo and estilo == "Heading 3":
            nivel3 += 1
            limpio = re.sub(r"^\d+\.\d+\.\d+\.?\s*(->\s*)?", "", texto)
            limpio = re.sub(r"^-\s*", "", limpio)
            p.text = f"{capitulo_actual}.{nivel2}.{nivel3}. {limpio}"


def convertir_listas_a_bullets(doc):
    """Convierte a viñetas (•) las listas que pandoc genera con abstractNum decimal
    o cuyo numId no está definido en la parte de numeración del documento combinado.

    El problema raíz: combinar_caratula_y_cuerpo copia el cuerpo XML de pandoc
    (con numPr referenciando numIds 58-68) al doc de carátula, pero NO copia la
    numbering_part. Así esos numIds quedan sin definición → Word los muestra como
    decimales. Esta función crea una definición bullet (•, Symbol) para todos ellos.
    """
    NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

    # Obtenemos la parte de numeración
    try:
        np_part = doc.part.numbering_part
        numbering_el = np_part._element
    except Exception:
        return  # Sin parte de numeración, nada que hacer

    # --- Paso 1: catalogar qué numIds usan párrafos "normales" (no encabezados)
    numids_heading: set[str] = set()
    numids_normal: set[str] = set()

    for p in doc.paragraphs:
        pPr = p._p.find(qn("w:pPr"))
        if pPr is None:
            continue
        numPr = pPr.find(qn("w:numPr"))
        if numPr is None:
            continue
        nid_el = numPr.find(qn("w:numId"))
        if nid_el is None:
            continue
        numId_val = nid_el.get(qn("w:val"))
        if not numId_val or numId_val == "0":
            continue
        sn = p.style.name if p.style else ""
        if sn.startswith("Heading"):
            numids_heading.add(numId_val)
        else:
            numids_normal.add(numId_val)

    # numIds usados sólo por párrafos normales (candidatos a convertir)
    candidatos = numids_normal - numids_heading

    if not candidatos:
        return

    # --- Paso 2: separar en definidos vs indefinidos ---
    definidos_decimales: set[str] = set()   # numIds con abstractNum decimal
    indefinidos: set[str] = set()           # numIds sin abstractNum en el doc

    def get_abstractNum(numId_str):
        num_el = numbering_el.find(f'{{{NS}}}num[@{{{NS}}}numId="{numId_str}"]')
        if num_el is None:
            return None
        abref = num_el.find(f'{{{NS}}}abstractNumId')
        if abref is None:
            return None
        abid = abref.get(f'{{{NS}}}val')
        return numbering_el.find(f'{{{NS}}}abstractNum[@{{{NS}}}abstractNumId="{abid}"]')

    for numId_str in candidatos:
        abn = get_abstractNum(numId_str)
        if abn is None:
            indefinidos.add(numId_str)
        else:
            # Check if any level is decimal
            for lvl in abn.findall(f'{{{NS}}}lvl'):
                nf = lvl.find(f'{{{NS}}}numFmt')
                if nf is not None and nf.get(f'{{{NS}}}val') == 'decimal':
                    definidos_decimales.add(numId_str)
                    break

    # --- Paso 3a: convertir en-place los abstractNums decimales ya definidos ---
    for numId_str in definidos_decimales:
        abn = get_abstractNum(numId_str)
        if abn is None:
            continue
        for lvl in abn.findall(f'{{{NS}}}lvl'):
            nf = lvl.find(f'{{{NS}}}numFmt')
            if nf is None or nf.get(f'{{{NS}}}val') != 'decimal':
                continue
            nf.set(f'{{{NS}}}val', 'bullet')
            lt = lvl.find(f'{{{NS}}}lvlText')
            if lt is None:
                lt = OxmlElement('w:lvlText')
                lvl.append(lt)
            lt.set(f'{{{NS}}}val', '\uf0b7')
            rpr = lvl.find(f'{{{NS}}}rPr')
            if rpr is None:
                rpr = OxmlElement('w:rPr')
                lvl.append(rpr)
            rfonts = rpr.find(f'{{{NS}}}rFonts')
            if rfonts is None:
                rfonts = OxmlElement('w:rFonts')
                rpr.insert(0, rfonts)
            rfonts.set(f'{{{NS}}}ascii', 'Symbol')
            rfonts.set(f'{{{NS}}}hAnsi', 'Symbol')
            rfonts.set(f'{{{NS}}}hint', 'default')


    # --- Paso 3b: crear UN único abstractNum bullet para todos los numIds sin definición ---
    if not indefinidos:
        return

    # Determinar un abstractNumId libre (mayor que los existentes)
    abstract_ids_existentes = [
        int(el.get(f'{{{NS}}}abstractNumId', '0'))
        for el in numbering_el.findall(f'{{{NS}}}abstractNum')
    ]
    nuevo_abstract_id = str(max(abstract_ids_existentes, default=0) + 1)

    # Construir el abstractNum bullet
    new_abn = OxmlElement('w:abstractNum')
    new_abn.set(f'{{{NS}}}abstractNumId', nuevo_abstract_id)

    # Solo nivel 0 — bullet
    new_lvl = OxmlElement('w:lvl')
    new_lvl.set(f'{{{NS}}}ilvl', '0')

    el_start = OxmlElement('w:start')
    el_start.set(f'{{{NS}}}val', '1')

    el_numFmt = OxmlElement('w:numFmt')
    el_numFmt.set(f'{{{NS}}}val', 'bullet')

    el_lvlText = OxmlElement('w:lvlText')
    el_lvlText.set(f'{{{NS}}}val', '\uf0b7')

    el_lvlJc = OxmlElement('w:lvlJc')
    el_lvlJc.set(f'{{{NS}}}val', 'left')

    el_pPr = OxmlElement('w:pPr')
    el_ind = OxmlElement('w:ind')
    el_ind.set(f'{{{NS}}}left', '720')
    el_ind.set(f'{{{NS}}}hanging', '360')
    el_pPr.append(el_ind)

    el_rPr = OxmlElement('w:rPr')
    el_rFonts = OxmlElement('w:rFonts')
    el_rFonts.set(f'{{{NS}}}ascii', 'Symbol')
    el_rFonts.set(f'{{{NS}}}hAnsi', 'Symbol')
    el_rFonts.set(f'{{{NS}}}hint', 'default')
    el_rPr.append(el_rFonts)

    new_lvl.append(el_start)
    new_lvl.append(el_numFmt)
    new_lvl.append(el_lvlText)
    new_lvl.append(el_lvlJc)
    new_lvl.append(el_pPr)
    new_lvl.append(el_rPr)
    new_abn.append(new_lvl)

    # Insertar abstractNum antes del primer w:num
    first_num = numbering_el.find(f'{{{NS}}}num')
    if first_num is not None:
        numbering_el.insert(list(numbering_el).index(first_num), new_abn)
    else:
        numbering_el.append(new_abn)

    # Determinar un numId libre para este abstractNum
    num_ids_existentes = [
        int(el.get(f'{{{NS}}}numId', '0'))
        for el in numbering_el.findall(f'{{{NS}}}num')
    ]
    nuevo_num_id = str(max(num_ids_existentes, default=0) + 1)

    # Crear el w:num que referencia el nuevo abstractNum
    new_num = OxmlElement('w:num')
    new_num.set(f'{{{NS}}}numId', nuevo_num_id)
    el_abId = OxmlElement('w:abstractNumId')
    el_abId.set(f'{{{NS}}}val', nuevo_abstract_id)
    new_num.append(el_abId)
    numbering_el.append(new_num)

    # Reasignar todos los párrafos con numIds indefinidos al nuevo numId
    for p in doc.paragraphs:
        pPr = p._p.find(qn("w:pPr"))
        if pPr is None:
            continue
        numPr = pPr.find(qn("w:numPr"))
        if numPr is None:
            continue
        nid_el = numPr.find(qn("w:numId"))
        if nid_el is None:
            continue
        if nid_el.get(qn("w:val")) in indefinidos:
            nid_el.set(qn("w:val"), nuevo_num_id)


def _ajustar_fuente_runs(parrafo, tamano_pt, negrita=None):
    """Ajusta tamaño (y opcionalmente negrita) en runs no vacíos."""
    for run in parrafo.runs:
        if run.text and run.text.strip():
            run.font.size = Pt(tamano_pt)
            if negrita is not None:
                run.bold = negrita


def _eliminar_parrafo(parrafo):
    """Elimina un párrafo del documento."""
    elemento = parrafo._element  # pyright: ignore[reportPrivateUsage]
    padre = elemento.getparent()
    if padre is not None:
        padre.remove(elemento)


def ajustar_caratula(doc):
    """Aproxima la carátula al formato esperado (logo, jerarquía y espaciado)."""
    # Aumentar tamaño del logo para que tenga mayor presencia en portada.
    if doc.inline_shapes:
        doc.inline_shapes[0].width = Inches(3.3)
        doc.inline_shapes[0].height = Inches(1.32)

    # Ubicar fin de carátula usando el año.
    fin = None
    for idx, p in enumerate(doc.paragraphs):
        if p.text.strip() == "2026":
            fin = idx
            break
    if fin is None:
        return

    # Limpiar líneas vacías de la portada (excepto logo y saltos explícitos con nbsp).
    for p in list(doc.paragraphs[: fin + 1]):
        texto = p.text.strip()
        tiene_nbsp = "\xa0" in p.text
        tiene_logo = "w:drawing" in p._p.xml
        if not texto and not tiene_logo and not tiene_nbsp:
            _eliminar_parrafo(p)

    # Recalcular fin de portada tras limpieza.
    fin = None
    for idx, p in enumerate(doc.paragraphs):
        if p.text.strip() == "2026":
            fin = idx
            break
    if fin is None:
        return

    # Espaciado adaptativo: si la carátula creciera demasiado, compacta.
    texto_portada = " ".join(p.text.strip() for p in doc.paragraphs[: fin + 1] if p.text.strip())
    modo_compacto = len(texto_portada) > 650
    factor = 0.82 if modo_compacto else 1.0

    for idx, p in enumerate(doc.paragraphs[: fin + 1]):
        texto = p.text.strip()
        tiene_nbsp = "\xa0" in p.text
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 1.15

        # Párrafo del logo (sin texto, con drawing embebido).
        if not texto and "w:drawing" in p._p.xml:
            p.paragraph_format.space_before = Pt(18 * factor)
            p.paragraph_format.space_after = Pt(22 * factor)
            continue

        # Saltos de línea explícitos colocados en el markdown de portada.
        if not texto and tiene_nbsp:
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(18 * factor)
            continue

        if texto == "MAESTRÍA EN INGENIERÍA DE SOFTWARE":
            _ajustar_fuente_runs(p, 17, True)
            p.paragraph_format.space_after = Pt(14 * factor)
        elif texto == "PROYECTO DE TESIS":
            _ajustar_fuente_runs(p, 16, True)
            p.paragraph_format.space_after = Pt(18 * factor)
        elif "MODELO DE INTEROPERABILIDAD BASADO EN HL7 FHIR" in texto:
            _ajustar_fuente_runs(p, 18, True)
            p.paragraph_format.space_after = Pt(18 * factor)
        elif texto == "PRESENTADO POR:":
            _ajustar_fuente_runs(p, 14, False)
            p.paragraph_format.space_after = Pt(12 * factor)
        elif texto.startswith("Bach."):
            _ajustar_fuente_runs(p, 14, True)
            p.paragraph_format.space_after = Pt(7 * factor)
        elif texto == "PARA OPTAR EL GRADO ACADÉMICO DE:":
            _ajustar_fuente_runs(p, 14, False)
            p.paragraph_format.space_before = Pt(10 * factor)
            p.paragraph_format.space_after = Pt(12 * factor)
        elif texto == "MAESTRO(A) EN INGENIERÍA DE SOFTWARE":
            _ajustar_fuente_runs(p, 16, True)
            p.paragraph_format.space_after = Pt(10 * factor)
        elif texto == "LIMA":
            _ajustar_fuente_runs(p, 13, True)
            p.paragraph_format.space_after = Pt(6 * factor)
        elif texto == "2026":
            _ajustar_fuente_runs(p, 13, True)


def excluir_preliminares_del_toc(doc):
    """Excluye secciones preliminares del TOC si es necesario."""
    excluir: set[str] = set()  # no hay secciones preliminares a excluir actualmente
    for p in doc.paragraphs:
        if p.text.strip() in excluir:
            p.style = doc.styles["Normal"]
            _ajustar_fuente_runs(p, 16, True)


def insertar_indices_automaticos(doc):
    """Inserta TOC y listas automáticas en sus respectivas secciones."""
    for p in list(doc.paragraphs):
        titulo = p.text.strip()
        if titulo == "Índice":
            destino = insertar_parrafo_despues(p)
            insertar_campo_word(destino, 'TOC \\o "1-3" \\h \\z \\u')
        elif titulo == "Índice de tablas":
            destino = insertar_parrafo_despues(p)
            insertar_campo_word(destino, 'TOC \\h \\z \\c "Table"')
        elif titulo == "Índice de figuras":
            destino = insertar_parrafo_despues(p)
            insertar_campo_word(destino, 'TOC \\h \\z \\c "Figure"')


def aplicar_bordes_tablas(doc):
    """Fuerza bordes visibles en todas las tablas."""
    for tabla in doc.tables:
        tabla.style = "Table Grid"

        tbl = tabla._tbl  # pyright: ignore[reportPrivateUsage]
        tbl_pr = tbl.tblPr

        bordes = tbl_pr.find(qn("w:tblBorders"))
        if bordes is not None:
            tbl_pr.remove(bordes)

        bordes = OxmlElement("w:tblBorders")
        for borde in ("top", "left", "bottom", "right", "insideH", "insideV"):
            el = OxmlElement(f"w:{borde}")
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), "8")
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), "000000")
            bordes.append(el)

        tbl_pr.append(bordes)


def actualizar_campos_con_word(destino):
    """Actualiza TOC/índices y campos usando Word COM para guardar números finales."""
    path = str(destino).replace("'", "''")
    script = rf"""
$ErrorActionPreference = 'Stop'
$path = '{path}'
$word = New-Object -ComObject Word.Application
$word.Visible = $false
$word.DisplayAlerts = 0
$doc = $word.Documents.Open($path)
$doc.Repaginate() | Out-Null
$doc.Fields.Update() | Out-Null
$doc.Repaginate() | Out-Null
foreach ($toc in $doc.TablesOfContents) {{ $toc.Update() }}
foreach ($tof in $doc.TablesOfFigures) {{ $tof.Update() }}
$doc.Fields.Update() | Out-Null
$doc.Save()
$doc.Close()
$word.Quit()
"""
    resultado = subprocess.run(
        ["powershell", "-NoProfile", "-ExecutionPolicy", "Bypass", "-Command", script],
        capture_output=True,
        text=True,
        check=False,
    )
    if resultado.returncode == 0:
        print("Campos de Word actualizados automáticamente (TOC/índices).")
    else:
        print("No se pudo actualizar campos con Word automáticamente.")
        if resultado.stderr:
            print(resultado.stderr)


def _pstyle_val(p):
    """Devuelve el w:val del pStyle leyendo directamente el XML (no depende de styles.xml)."""
    ppr = p._p.find(qn("w:pPr"))  # pyright: ignore[reportPrivateUsage]
    if ppr is None:
        return ""
    ps = ppr.find(qn("w:pStyle"))
    return ps.get(qn("w:val"), "") if ps is not None else ""


def numerar_captions_figuras(doc):
    """Agrega campo SEQ Figure a captions de imagen para que el índice de figuras funcione."""
    for p in doc.paragraphs:
        if _pstyle_val(p) != "ImageCaption":
            continue
        texto_original = p.text.strip()
        if not texto_original:
            continue

        # Limpiar todos los runs existentes (preservar pPr)
        p_elem = p._p  # pyright: ignore[reportPrivateUsage]
        for child in list(p_elem):
            if child.tag != qn("w:pPr"):
                p_elem.remove(child)

        # Run: "Figura "
        run_pre = p.add_run("Figura ")
        run_pre.font.size = Pt(10)
        run_pre.font.name = "Arial"
        run_pre.italic = True

        # Campo SEQ Figure — cada componente en su propio run (requerido por Word)
        r_begin = p.add_run()._r  # pyright: ignore[reportPrivateUsage]
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        r_begin.append(fld_begin)

        r_instr = p.add_run()._r  # pyright: ignore[reportPrivateUsage]
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = " SEQ Figure \\* ARABIC "
        r_instr.append(instr)

        r_sep = p.add_run()._r  # pyright: ignore[reportPrivateUsage]
        fld_sep = OxmlElement("w:fldChar")
        fld_sep.set(qn("w:fldCharType"), "separate")
        r_sep.append(fld_sep)

        run_num = p.add_run("#")
        run_num.font.size = Pt(10)
        run_num.font.name = "Arial"
        run_num.italic = True

        r_end = p.add_run()._r  # pyright: ignore[reportPrivateUsage]
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        r_end.append(fld_end)

        # Run: ". " + texto original del caption
        run_post = p.add_run(f". {texto_original}")
        run_post.font.size = Pt(10)
        run_post.font.name = "Arial"
        run_post.italic = True

        # Centrar el caption
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def numerar_captions_tablas(doc):
    """Agrega campo SEQ Table a captions de tabla para que el índice de tablas funcione."""
    # Detectar captions por posición: párrafo directamente antes de w:tbl en el body
    body = doc.element.body
    p_ids_caption: set[int] = set()
    children = list(body)
    for i, child in enumerate(children):
        if child.tag == qn("w:tbl") and i > 0:
            prev = children[i - 1]
            if prev.tag == qn("w:p"):
                p_ids_caption.add(id(prev))

    for p in doc.paragraphs:
        pstyle = _pstyle_val(p)
        es_caption = (
            id(p._p) in p_ids_caption  # pyright: ignore[reportPrivateUsage]
            or pstyle in {"TableCaption", "TableTitle", "Tabletitre"}
        )
        if not es_caption:
            continue

        texto_original = re.sub(r"\s*\{#[^}]+\}", "", p.text.strip()).strip()
        # Quitar prefijo "Tabla N. " o "Tabla N: " que pandoc ya incluyó en el texto
        texto_original = re.sub(r"^Tabla\s+\d+\s*[\.:]\.?\s*", "", texto_original).strip()
        if not texto_original or texto_original.startswith("Nota."):
            continue

        # Limpiar todos los runs existentes (preservar pPr)
        p_elem = p._p  # pyright: ignore[reportPrivateUsage]
        for child in list(p_elem):
            if child.tag != qn("w:pPr"):
                p_elem.remove(child)

        # Run: "Tabla "
        run_pre = p.add_run("Tabla ")
        run_pre.font.size = Pt(10)
        run_pre.font.name = "Arial"

        # Campo SEQ Table — cada componente en su propio run (requerido por Word)
        r_begin = p.add_run()._r  # pyright: ignore[reportPrivateUsage]
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        r_begin.append(fld_begin)

        r_instr = p.add_run()._r  # pyright: ignore[reportPrivateUsage]
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = " SEQ Table \\* ARABIC "
        r_instr.append(instr)

        r_sep = p.add_run()._r  # pyright: ignore[reportPrivateUsage]
        fld_sep = OxmlElement("w:fldChar")
        fld_sep.set(qn("w:fldCharType"), "separate")
        r_sep.append(fld_sep)

        run_num = p.add_run("#")
        run_num.font.size = Pt(10)
        run_num.font.name = "Arial"

        r_end = p.add_run()._r  # pyright: ignore[reportPrivateUsage]
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        r_end.append(fld_end)

        # Run: ". " + texto original del caption
        run_post = p.add_run(f". {texto_original}")
        run_post.font.size = Pt(10)
        run_post.font.name = "Arial"


def formatear_notas_fuente(doc):
    """Centra y aplica Arial 10pt a TODAS las notas ('Nota. ...') fuera de tablas."""
    for p in doc.paragraphs:
        texto = p.text.strip()
        if not texto.startswith("Nota."):
            continue
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_fuente_parrafo_ppr(p, 10, "Arial")
        for run in p.runs:
            _set_fuente_run(run, 10, "Arial")


# ---------------------------------------------------------------------------
# Secciones landscape automáticas para tablas anchas
# ---------------------------------------------------------------------------

def _contar_columnas_tabla(tbl_elem):
    """Cuenta columnas de una tabla DOCX usando w:tblGrid o la primera fila."""
    tbl_grid = tbl_elem.find(qn("w:tblGrid"))
    if tbl_grid is not None:
        return len(tbl_grid.findall(qn("w:gridCol")))
    first_row = tbl_elem.find(qn("w:tr"))
    if first_row is not None:
        return len(first_row.findall(qn("w:tc")))
    return 0


def _texto_elem(elem):
    """Texto concatenado de todos los w:t dentro del elemento."""
    return "".join(t.text or "" for t in elem.iter(qn("w:t"))).strip()


def _es_caption_tabla(elem):
    """True si el elemento es un párrafo caption de tabla (por estilo o texto)."""
    if elem.tag != qn("w:p"):
        return False
    ppr = elem.find(qn("w:pPr"))
    if ppr is not None:
        ps = ppr.find(qn("w:pStyle"))
        if ps is not None and ps.get(qn("w:val"), "") in {"TableCaption", "TableTitle", "Tabletitre"}:
            return True
    # Fallback: comienza con "Tabla" (tras numerar_captions_tablas)
    return _texto_elem(elem).startswith("Tabla")


def _hacer_sectpr(landscape: bool):
    """Crea w:sectPr A4 en orientación vertical (portrait) u horizontal (landscape)."""
    sect = OxmlElement("w:sectPr")
    pgSz = OxmlElement("w:pgSz")
    if landscape:
        pgSz.set(qn("w:w"), "16838")   # 29.7 cm → twips
        pgSz.set(qn("w:h"), "11906")   # 21 cm  → twips
        pgSz.set(qn("w:orient"), "landscape")
    else:
        pgSz.set(qn("w:w"), "11906")
        pgSz.set(qn("w:h"), "16838")
    sect.append(pgSz)
    pgMar = OxmlElement("w:pgMar")
    # Márgenes 1 pulgada (1440 twips) en landscape, reducir laterales para más espacio
    if landscape:
        pgMar.set(qn("w:top"), "720")
        pgMar.set(qn("w:right"), "720")
        pgMar.set(qn("w:bottom"), "720")
        pgMar.set(qn("w:left"), "720")
    else:
        pgMar.set(qn("w:top"), "1440")
        pgMar.set(qn("w:right"), "1440")
        pgMar.set(qn("w:bottom"), "1440")
        pgMar.set(qn("w:left"), "1440")
    pgMar.set(qn("w:header"), "708")
    pgMar.set(qn("w:footer"), "708")
    pgMar.set(qn("w:gutter"), "0")
    sect.append(pgMar)
    return sect


def _parrafo_seccion(landscape: bool):
    """Párrafo vacío cuyo pPr contiene un sectPr para marcar inicio de sección."""
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    pPr.append(_hacer_sectpr(landscape))
    p.append(pPr)
    return p


def aplicar_orientacion_tablas_anchas(doc, min_columnas: int = 4):
    """Envuelve tablas de ≥ min_columnas columnas en secciones landscape.

    Dos tablas anchas consecutivas (separadas solo por párrafos vacíos o notas)
    se mantienen en la misma sección landscape.
    """
    body = doc.element.body

    # Elementos del body sin el w:sectPr final
    ch = [c for c in body if c.tag != qn("w:sectPr")]
    n = len(ch)

    # ── Paso 1: encontrar bloques (start, end) ──────────────────────────────
    bloques: list[tuple[int, int]] = []
    i = 0
    while i < n:
        if ch[i].tag != qn("w:tbl") or _contar_columnas_tabla(ch[i]) < min_columnas:
            i += 1
            continue

        # Tabla ancha encontrada en ch[i]
        start = i
        # Extender hacia atrás: incluir caption si el párrafo anterior lo es
        if start > 0 and _es_caption_tabla(ch[start - 1]):
            start -= 1

        end = i
        # Extender hacia adelante: incluir párrafos de nota ("Nota. ...")
        j = end + 1
        while j < n and ch[j].tag == qn("w:p") and _texto_elem(ch[j]).startswith("Nota."):
            end = j
            j += 1

        bloques.append((start, end))
        i = end + 1

    if not bloques:
        return

    # ── Paso 2: fusionar bloques separados solo por párrafos vacíos ─────────
    fusionados: list[tuple[int, int]] = [bloques[0]]
    for blk_s, blk_e in bloques[1:]:
        prev_s, prev_e = fusionados[-1]
        gap_vacios = all(
            ch[k].tag == qn("w:p") and not _texto_elem(ch[k])
            for k in range(prev_e + 1, blk_s)
        ) if blk_s > prev_e + 1 else True
        if gap_vacios:
            fusionados[-1] = (prev_s, blk_e)
        else:
            fusionados.append((blk_s, blk_e))

    # ── Paso 3: capturar referencias antes de modificar el body ─────────────
    grupos_elems = [(ch[s], ch[e]) for s, e in fusionados]

    # ── Paso 4: insertar section breaks en orden inverso ────────────────────
    for first_elem, last_elem in reversed(grupos_elems):
        # Párrafo con sectPr(landscape) DESPUÉS del bloque → cierra sección landscape
        last_elem.addnext(_parrafo_seccion(landscape=True))
        # Párrafo con sectPr(portrait) ANTES del bloque → cierra sección portrait anterior
        first_elem.addprevious(_parrafo_seccion(landscape=False))

    tablas_envueltas = len(grupos_elems)
    print(f"  Secciones landscape aplicadas: {tablas_envueltas} grupo(s) de tablas anchas.")


# ---------------------------------------------------------------------------
# Ajuste dinámico de anchos de columnas proporcional al contenido
# ---------------------------------------------------------------------------

# Ancho de texto disponible en twips para cada orientación
_ANCHO_LANDSCAPE = 16838 - 1440   # A4 landscape - márgenes 720+720
_ANCHO_PORTRAIT  = 11906 - 2880   # A4 portrait  - márgenes 1440+1440
_MIN_COL_TWIPS   = 800             # ancho mínimo garantizado por columna (~1.4 cm)


def _max_len_por_columna(tbl_elem):
    """Longitud máxima de texto (caracteres) por columna, manejando celdas fusionadas."""
    tbl_grid = tbl_elem.find(qn("w:tblGrid"))
    n_cols = len(tbl_grid.findall(qn("w:gridCol"))) if tbl_grid is not None else 0
    if n_cols == 0:
        first_row = tbl_elem.find(qn("w:tr"))
        if first_row is None:
            return []
        n_cols = len(first_row.findall(qn("w:tc")))
    if n_cols == 0:
        return []

    max_len: list[int] = [0] * n_cols
    for tr in tbl_elem.findall(qn("w:tr")):
        col_idx = 0
        for tc in tr.findall(qn("w:tc")):
            if col_idx >= n_cols:
                break
            tc_pr = tc.find(qn("w:tcPr"))
            span = 1
            if tc_pr is not None:
                gs = tc_pr.find(qn("w:gridSpan"))
                if gs is not None:
                    span = int(gs.get(qn("w:val"), "1") or "1")
            texto = "".join(t.text or "" for t in tc.iter(qn("w:t"))).strip()
            # Imputar solo a la primera columna del span (heurística conservadora)
            max_len[col_idx] = max(max_len[col_idx], len(texto))
            col_idx += span
    return max_len


def _distribuir_anchos(max_len: list[int], ancho_total: int) -> list[int]:
    """Distribuye ancho_total (twips) entre columnas de forma proporcional.

    Usa raíz cuadrada de la longitud para suavizar el impacto de columnas
    con mucho texto y evitar que dominen excesivamente el espacio.
    """
    n = len(max_len)
    if n == 0:
        return []
    # Raíz cuadrada para distribución suavizada
    pesos = [math.sqrt(max(ml, 1)) for ml in max_len]
    total_peso = sum(pesos) or 1
    # Ancho proporcional con piso mínimo
    espacio_libre = max(0, ancho_total - _MIN_COL_TWIPS * n)
    anchos = [
        _MIN_COL_TWIPS + int(espacio_libre * (p / total_peso))
        for p in pesos
    ]
    # Ajustar diferencia de redondeo en la última columna
    anchos[-1] += ancho_total - sum(anchos)
    return anchos


def _aplicar_anchos_a_tabla(tbl_elem, anchos: list[int]):
    """Escribe los anchos calculados en w:tblGrid, w:tblW y cada celda."""
    n = len(anchos)
    ancho_total = sum(anchos)

    # w:tblGrid
    tbl_grid = tbl_elem.find(qn("w:tblGrid"))
    if tbl_grid is not None:
        for k, gc in enumerate(tbl_grid.findall(qn("w:gridCol"))):
            if k < n:
                gc.set(qn("w:w"), str(anchos[k]))

    # w:tblPr → w:tblW y w:tblLayout
    tbl_pr = tbl_elem.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(ancho_total))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    # Ancho de cada celda
    for tr in tbl_elem.findall(qn("w:tr")):
        col_idx = 0
        for tc in tr.findall(qn("w:tc")):
            if col_idx >= n:
                break
            tc_pr = tc.find(qn("w:tcPr"))
            if tc_pr is None:
                tc_pr = OxmlElement("w:tcPr")
                tc.insert(0, tc_pr)
            span = 1
            gs = tc_pr.find(qn("w:gridSpan"))
            if gs is not None:
                span = int(gs.get(qn("w:val"), "1") or "1")
            w_val = sum(anchos[col_idx: col_idx + span]) if col_idx + span <= n else anchos[col_idx]
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(w_val))
            tc_w.set(qn("w:type"), "dxa")
            col_idx += span


def _set_fuente_run(run, tamano_pt: float, fuente: str = "Arial"):
    """Fija fuente y tamaño en un run y sus marcas rPr XML."""
    run.font.name = fuente
    run.font.size = Pt(tamano_pt)
    r = run._r  # pyright: ignore[reportPrivateUsage]
    rPr = r.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        r.insert(0, rPr)
    for tag, attr, val in [
        ("w:sz",     qn("w:val"), str(int(tamano_pt * 2))),
        ("w:szCs",   qn("w:val"), str(int(tamano_pt * 2))),
        ("w:rFonts", qn("w:ascii"), fuente),
        ("w:rFonts", qn("w:hAnsi"), fuente),
        ("w:rFonts", qn("w:cs"),    fuente),
    ]:
        el = rPr.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            rPr.append(el)
        el.set(attr, val)
        if tag == "w:rFonts":
            el.set(qn("w:eastAsia"), fuente)


def _set_fuente_parrafo_ppr(p, tamano_pt: float, fuente: str = "Arial"):
    """Inyecta sz/szCs/rFonts en el w:pPr/w:rPr para cubrir runs vacíos y herencia."""
    p_elem = p._p  # pyright: ignore[reportPrivateUsage]
    pPr = p_elem.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        p_elem.insert(0, pPr)
    rPr_pPr = pPr.find(qn("w:rPr"))
    if rPr_pPr is None:
        rPr_pPr = OxmlElement("w:rPr")
        pPr.append(rPr_pPr)
    for tag, attr, val in [
        ("w:sz",     qn("w:val"), str(int(tamano_pt * 2))),
        ("w:szCs",   qn("w:val"), str(int(tamano_pt * 2))),
        ("w:rFonts", qn("w:ascii"), fuente),
        ("w:rFonts", qn("w:hAnsi"), fuente),
        ("w:rFonts", qn("w:cs"),    fuente),
    ]:
        el = rPr_pPr.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            rPr_pPr.append(el)
        el.set(attr, val)
        if tag == "w:rFonts":
            el.set(qn("w:eastAsia"), fuente)


def aplicar_fuente_arial_documento(doc, tamano_pt: float = 12, fuente: str = "Arial"):
    """Aplica Arial 12pt a todos los párrafos del cuerpo que no sean tabla ni caption."""
    estilos_excluidos = {"ImageCaption", "TableCaption", "TableTitle", "Tabletitre", "Caption"}

    for p in doc.paragraphs:
        estilo = _pstyle_val(p)
        if estilo in estilos_excluidos:
            continue
        # Exclusiones por texto: notas, captions figura, captions tabla
        texto = p.text.strip()
        if texto.startswith("Nota."):
            continue
        if re.match(r"^Figura\s", texto):
            continue
        if re.match(r"^Tabla\s+[\d#]", texto):
            continue
        _set_fuente_parrafo_ppr(p, tamano_pt, fuente)
        for run in p.runs:
            _set_fuente_run(run, tamano_pt, fuente)


def ajustar_fuente_tablas(doc, tamano_pt: float = 10, fuente: str = "Arial"):
    """Establece Arial 10pt en todo el contenido de las tablas."""
    for tabla in doc.tables:
        for row in tabla.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _set_fuente_parrafo_ppr(p, tamano_pt, fuente)
                    for run in p.runs:
                        _set_fuente_run(run, tamano_pt, fuente)


def ajustar_anchos_tablas(doc, umbral_landscape: int = 4):
    """Ajusta anchos de columnas proporcionales al contenido máximo.

    Tablas con ≥ umbral_landscape columnas reciben el ancho de página landscape;
    las demás, el ancho de página portrait.
    """
    ajustadas = 0
    for tabla in doc.tables:
        tbl_elem = tabla._tbl  # pyright: ignore[reportPrivateUsage]
        n_cols = _contar_columnas_tabla(tbl_elem)
        if n_cols < 2:
            continue
        ancho = _ANCHO_LANDSCAPE if n_cols >= umbral_landscape else _ANCHO_PORTRAIT
        max_len = _max_len_por_columna(tbl_elem)
        if not max_len:
            continue
        anchos = _distribuir_anchos(max_len, ancho)
        _aplicar_anchos_a_tabla(tbl_elem, anchos)
        ajustadas += 1
    print(f"  Anchos dinámicos aplicados: {ajustadas} tabla(s).")


def postprocesar_docx(destino):
    """Aplica formato y campos Word después de la conversión de pandoc."""
    doc = Document(str(destino))
    excluir_preliminares_del_toc(doc)
    aplicar_numeracion_titulos(doc)
    convertir_listas_a_bullets(doc)
    aplicar_saltos_pagina_encabezados(doc)
    numerar_captions_figuras(doc)
    numerar_captions_tablas(doc)
    formatear_notas_fuente(doc)
    insertar_indices_automaticos(doc)
    aplicar_bordes_tablas(doc)
    aplicar_orientacion_tablas_anchas(doc)
    ajustar_anchos_tablas(doc)
    aplicar_fuente_arial_documento(doc)
    ajustar_fuente_tablas(doc)
    activar_actualizacion_campos(doc)
    doc.save(str(destino))
    actualizar_campos_con_word(destino)
    print("Postproceso DOCX completado (índices, saltos y numeración).")


if __name__ == "__main__":
    print("=== Preparando reference doc con estilos ===")
    preparar_reference_doc()

    caratula_entrada = CARATULA_MANUAL if CARATULA_MANUAL.exists() else CARATULA_MANUAL_ALT
    if not caratula_entrada.exists():
        print(f"ERROR: No se encontró la carátula manual ({CARATULA_MANUAL.name} o {CARATULA_MANUAL_ALT.name})")
        sys.exit(1)

    md_cuerpo_tmp = extraer_markdown_cuerpo(ENTRADA)
    try:
        print("\n=== Ejecutando pandoc (solo cuerpo) ===")
        exito = ejecutar_pandoc(md_cuerpo_tmp, CUERPO_TEMP, REFERENCE)
        if not exito:
            sys.exit(1)

        print("\n=== Combinando carátula manual + cuerpo ===")
        salida_objetivo = SALIDA
        try:
            combinar_caratula_y_cuerpo(caratula_entrada, CUERPO_TEMP, salida_objetivo)
        except PermissionError:
            print("El archivo principal parece estar bloqueado. Se generará salida alternativa.")
            salida_objetivo = SALIDA_ALT
            combinar_caratula_y_cuerpo(caratula_entrada, CUERPO_TEMP, salida_objetivo)

        print("\n=== Aplicando postproceso DOCX ===")
        postprocesar_docx(salida_objetivo)
    finally:
        if CUERPO_TEMP.exists():
            CUERPO_TEMP.unlink()
        if md_cuerpo_tmp.exists():
            md_cuerpo_tmp.unlink()
